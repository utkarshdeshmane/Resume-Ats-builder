import os
import io
import uuid
import logging
from typing import Dict, List, Any
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Inches
import PyPDF2 as pdf
from dotenv import load_dotenv
import google.generativeai as genai

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Configure Gemini
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

genai.configure(api_key=GOOGLE_API_KEY)

# Constants
OUTPUTS_DIR = "outputs"
TEMPLATES_DIR = "templates"
RESUME_TEMPLATE = "resume_template.html"

# Section identifiers
SECTION_MARKERS = {
    "summary": ["summary", "professional summary", "overview"],
    "skills": ["skills", "technical skills", "competencies"],
    "experience": ["experience", "work experience", "employment history"],
    "education": ["education", "academic background", "qualifications"],
    "certifications": ["certifications", "certificates", "professional certifications"],
    "projects": ["projects", "personal projects", "portfolio"]
}

class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors"""
    pass

def setup_directories() -> None:
    """Create necessary directories if they don't exist"""
    Path(OUTPUTS_DIR).mkdir(exist_ok=True)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file
    """
    try:
        reader = pdf.PdfReader(io.BytesIO(file_bytes))
        text = " ".join(page.extract_text() or "" for page in reader.pages)
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise ResumeParsingError(f"Failed to extract text from PDF: {str(e)}")

def parse_gemini_resume_to_data(text: str) -> Dict[str, Any]:
    """
    Parse resume text into structured data
    """
    data = {
        "name": "",
        "email": "",
        "phone": "",
        "summary": "",
        "skills": [],
        "experience": [],
        "education": [],
        "certifications": [],
        "projects": []
    }

    current_section = None
    buffer = []

    lines = text.strip().split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Contact Info
        if ":" in line:
            key, value = line.split(":", 1)
            key = key.lower().strip()
            if key in ["name", "email", "phone"]:
                data[key] = value.strip()
                continue

        # Section Headers
        section_found = False
        for section, markers in SECTION_MARKERS.items():
            if any(marker in line.lower() for marker in markers):
                if current_section and buffer:
                    data[current_section] = _process_section_buffer(buffer, current_section)
                current_section = section
                buffer = []
                section_found = True
                break

        if not section_found and current_section:
            buffer.append(line)

    # Process last section
    if current_section and buffer:
        data[current_section] = _process_section_buffer(buffer, current_section)

    return data

def _process_section_buffer(buffer: List[str], section: str) -> Any:
    """Process accumulated lines for a section"""
    if section == "summary":
        return " ".join(buffer)
    else:
        return [line.lstrip("-* ").strip() for line in buffer if line]

def render_pdf(data: Dict[str, Any]) -> str:
    """
    Render resume data to PDF
    """
    try:
        env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))
        template = env.get_template(RESUME_TEMPLATE)
        html_out = template.render(data)

        filename = f"resume_{uuid.uuid4().hex}.pdf"
        pdf_path = Path(OUTPUTS_DIR) / filename

        HTML(string=html_out).write_pdf(str(pdf_path))
        return str(pdf_path)
    except Exception as e:
        logger.error(f"PDF rendering failed: {str(e)}")
        raise

def render_word_document(data: Dict[str, Any]) -> str:
    """
    Render resume data to a Word document
    """
    try:
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Contact Info
        doc.add_heading(data.get('name', 'Name'), 0)
        contact_info = doc.add_paragraph()
        contact_info.add_run(f"Email: {data.get('email', '')} | Phone: {data.get('phone', '')}")

        # Summary
        if data.get('summary'):
            doc.add_heading('Professional Summary', 1)
            doc.add_paragraph(data['summary'])

        # Skills
        if data.get('skills'):
            doc.add_heading('Skills', 1)
            skills_para = doc.add_paragraph()
            skills_para.add_run(' • '.join(data['skills']))

        # Experience
        if data.get('experience'):
            doc.add_heading('Professional Experience', 1)
            for exp in data['experience']:
                doc.add_paragraph(exp, style='List Bullet')

        # Education
        if data.get('education'):
            doc.add_heading('Education', 1)
            for edu in data['education']:
                doc.add_paragraph(edu, style='List Bullet')

        # Certifications
        if data.get('certifications'):
            doc.add_heading('Certifications', 1)
            for cert in data['certifications']:
                doc.add_paragraph(cert, style='List Bullet')

        # Projects
        if data.get('projects'):
            doc.add_heading('Projects', 1)
            for project in data['projects']:
                doc.add_paragraph(project, style='List Bullet')

        filename = f"resume_{uuid.uuid4().hex}.docx"
        file_path = Path(OUTPUTS_DIR) / filename
        doc.save(str(file_path))

        return str(file_path)

    except Exception as e:
        logger.error(f"Word document generation failed: {str(e)}")
        raise

def render_resume(data: Dict[str, Any], output_path: str) -> None:
    """
    Render resume data to either PDF or Word format based on output path extension
    """
    try:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if output_path.suffix.lower() == '.pdf':
            rendered_path = render_pdf(data)
            Path(rendered_path).rename(output_path)

        elif output_path.suffix.lower() == '.docx':
            rendered_path = render_word_document(data)
            Path(rendered_path).rename(output_path)

        else:
            raise ValueError(f"Unsupported output format: {output_path.suffix}")

    except Exception as e:
        logger.error(f"Failed to render resume: {str(e)}")
        raise Exception(f"Resume generation failed: {str(e)}")

async def get_gemini_response(prompt: str) -> str:
    """
    Get response from Google's Gemini API
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)

        if not response.text:
            raise Exception("Empty response from Gemini API")

        return response.text.strip()

    except Exception as e:
        logger.error(f"Gemini API call failed: {str(e)}")
        raise Exception(f"Failed to get AI response: {str(e)}")
